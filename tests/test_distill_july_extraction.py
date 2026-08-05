"""July-2026 handoff extraction: continuation-line values, generic section
markers, material cards, and subservice city/page-kind derivation.

These cover the field-drop defects fixed in pipeline.generate.distill after the
emitter (built on April's inline format) met July's structurally different
handoff:

  1. CONTINUATION-LINE VALUES — a label on one paragraph with its value on the
     NEXT ('Hero Subheading:' then a separate line). April's inline form must
     keep working unchanged (additive).
  2. GENERIC SECTION MARKERS — page-specific '* SECTION' headings the fixed hub
     vocabulary does not enumerate (VENTILATION / MATERIALS / HOA / ...) must
     reach the page instead of being silently dropped.
  3. MATERIAL CARDS — 'Card N: Title' + body paragraphs -> a types grid.
  4. SUBSERVICE identity — a 3-segment slug is page_kind 'subservice' and the
     city comes from the slug, not from the meta-title page name.

All hermetic: DOCX synthesised with python-docx under tmp_path.
"""
from __future__ import annotations

import docx
import pytest

from pipeline.generate.distill import (
    _city_from_slug,
    build_draft,
    distill,
    label_value,
    read_docx_blocks,
    segment_pages,
)
from pipeline.generate.models import Section


def _save(document, tmp_path, name: str) -> str:
    path = tmp_path / name
    document.save(str(path))
    return str(path)


def _full_july_page(tmp_path) -> str:
    """A representative July subservice page: Page-Title boundary, a Hero
    Subheading on a CONTINUATION line, an INTRO, a generic VENTILATION prose
    section, a generic MATERIALS card section ('Card N: Title' + body), and an
    FAQ. Mirrors the real Acme July handoff shape."""
    d = docx.Document()
    d.add_paragraph('Page Title (56 chars): Soffit Installation in Matthews, NC | Since 2000')
    d.add_paragraph('Meta Description: Professional soffit installation in Matthews NC with '
                    'vented and solid options, code-compliant ventilation and free estimates.')
    d.add_paragraph('Canonical URL: https://acmeroofing.example.com/charlotte-nc/matthews/soffit-installation/')
    d.add_paragraph('HERO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Badge: Serving Matthews, NC')
    d.add_paragraph('H1: Soffit Installation in Matthews, NC')
    d.add_paragraph('Hero Subheading:')                              # label, no inline value
    d.add_paragraph('Soffit protects your roof from moisture and pests, and Acme installs it '
                    'right the first time for every Matthews home we serve.')  # value on next line
    d.add_paragraph('INTRO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Section Label: Trusted Since 2000')
    d.add_paragraph('H2: Soffit Installation Built for Matthews Homes')
    d.add_paragraph('Paragraph 1: Soffit installation in Matthews typically involves repairing '
                    'damaged areas or upgrading older homes with vented panels for airflow.')
    d.add_paragraph('VENTILATION SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Section Label: Why Attic Ventilation Matters')
    d.add_paragraph('H2: Why Correct Soffit Ventilation Is Critical on Matthews Homes')
    d.add_paragraph('Mecklenburg County code requires a balanced intake and exhaust ratio, and '
                    'older solid soffit provides zero intake, forcing the upper vents to carry '
                    'the whole load until the system fails.')                 # bare body prose
    d.add_paragraph('MATERIALS SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Section Label: Soffit Materials')
    d.add_paragraph('H2: Soffit Materials We Install in Matthews')
    d.add_paragraph('Card 1: Vinyl Soffit')                          # colon card head
    d.add_paragraph('Vinyl is rot-resistant, needs no painting, and withstands humidity and '
                    'pests, which makes it the economical default for most Matthews homes.')
    d.add_paragraph('Card 2: Aluminum Soffit')
    d.add_paragraph('Aluminum resists dents and cracking better than vinyl and suits lots with '
                    'heavy tree canopy where falling debris is a constant risk to the eave.')
    d.add_paragraph('Card 3: Fiber Cement Soffit')
    d.add_paragraph('Fiber cement is impervious to moisture, rot, and pests and holds paint far '
                    'longer than wood, which makes it the premium HOA-compliant choice here.')
    d.add_paragraph('FAQ SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Section Label: FAQs')
    d.add_paragraph('H2: Frequently Asked Questions')
    d.add_paragraph('Q1: Does my Matthews home need vented or solid soffit?')
    d.add_paragraph('Most homes need a mix to meet the ventilation ratio required by county '
                    'code, and Acme calculates the exact net free area needed for your attic '
                    'size before any material is ordered so the install is right.')
    return _save(d, tmp_path, 'july_full.docx')


# ---------------------------------------------------------------------------
# 1. Continuation-line value association  (the named fix)
# ---------------------------------------------------------------------------

def test_label_value_reads_continuation_line(tmp_path):
    """A bare 'Hero Subheading:' followed by its value on the next paragraph
    resolves via label_value to that following text."""
    d = docx.Document()
    d.add_paragraph('HERO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Hero Subheading:')
    d.add_paragraph('The value lives on the paragraph after the label.')
    d.add_paragraph('Trust Stats:')  # next label — the continuation must stop here
    blocks = read_docx_blocks(str(_save(d, tmp_path, 'cont.docx')))
    anchor = next(b for b in blocks if b.label.lower() == 'hero subheading')
    assert anchor.value == ''  # no inline value in the source
    assert label_value(blocks, anchor) == 'The value lives on the paragraph after the label.'


def test_label_value_inline_form_is_untouched(tmp_path):
    """April's inline 'Hero Subheading: <text>' resolves to exactly the inline
    value — the continuation branch never fires (additive guarantee)."""
    d = docx.Document()
    d.add_paragraph('HERO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Hero Subheading: Inline value on the same line as the label.')
    d.add_paragraph('The following paragraph must NOT be swallowed into the value.')
    blocks = read_docx_blocks(str(_save(d, tmp_path, 'inline.docx')))
    anchor = next(b for b in blocks if b.label.lower() == 'hero subheading')
    assert anchor.value == 'Inline value on the same line as the label.'
    assert label_value(blocks, anchor) == 'Inline value on the same line as the label.'


def test_july_hero_description_is_populated_not_empty(tmp_path):
    """The end-to-end symptom: the hero description must carry the continuation
    value, not read empty (which produced the false no_hero_copy block)."""
    draft = distill(_full_july_page(tmp_path))[0]
    assert draft.hero.description.strip()
    assert 'moisture and pests' in draft.hero.description
    assert not any(f.code == 'no_hero_copy' for f in draft.findings)


# ---------------------------------------------------------------------------
# 2 & 3. Generic section markers + material cards
# ---------------------------------------------------------------------------

def test_generic_prose_section_becomes_content_block(tmp_path):
    """An unrecognized 'VENTILATION SECTION' with bare body prose is extracted as
    a content-block, not silently dropped."""
    draft = distill(_full_july_page(tmp_path))[0]
    vent = [s for s in draft.real_sections()
            if isinstance(s, Section) and 'Ventilation Is Critical' in str(s.props.get('title', ''))]
    assert len(vent) == 1
    assert vent[0].type == 'content-block'
    assert any('intake' in p for p in vent[0].props['content'])


def test_generic_card_section_becomes_types_grid(tmp_path):
    """MATERIALS 'Card N: Title' + body paragraphs -> a types grid with each
    card's title and description populated."""
    draft = distill(_full_july_page(tmp_path))[0]
    mats = [s for s in draft.real_sections()
            if isinstance(s, Section) and s.type == 'types'
            and 'Materials' in str(s.props.get('title', ''))]
    assert len(mats) == 1
    cards = mats[0].props['cards']
    titles = [c['title'] for c in cards]
    assert titles == ['Vinyl Soffit', 'Aluminum Soffit', 'Fiber Cement Soffit']
    assert all(c['description'].strip() for c in cards)          # nothing empty
    assert 'rot-resistant' in cards[0]['description']


def test_no_silent_section_drop(tmp_path):
    """Core-body words must reflect the extracted sections — the July page is not
    a hollow shell (regression guard for the dropped-marker defect)."""
    draft = distill(_full_july_page(tmp_path))[0]
    # Intro alone is ~20 words; the generic VENTILATION + MATERIALS sections that
    # were being dropped add the rest. Comfortably above intro-only proves they landed.
    assert draft.core_body_words >= 100


# ---------------------------------------------------------------------------
# 4. Subservice identity — page_kind + city from slug
# ---------------------------------------------------------------------------

def test_subservice_page_kind_and_city_from_slug(tmp_path):
    """A 3-segment slug is page_kind 'subservice' and the city is 'Matthews'
    (from the slug) — never the full meta-title page name."""
    draft = distill(_full_july_page(tmp_path))[0]
    assert draft.page_kind == 'subservice'
    assert draft.city == 'Matthews'
    assert draft.state == 'NC'
    assert draft.slug == 'charlotte-nc/matthews/soffit-installation'
    # the full meta-title page name (its pipe-joined tail) never leaked into any
    # section label / title (the city='<meta title>' bug polluted these).
    leak = 'Matthews, NC | Since 2000'
    assert not any(leak in str(v)
                   for s in draft.real_sections() for v in s.props.values()
                   if isinstance(v, str))


# ── BUG-016 residual: _city_from_slug hardcoded Acme's states ───────────────
# The BUG-016 fix generalised the `leaf` path to read states_served, but
# _city_from_slug() kept the literal nc|sc|va|wv and never received config, so an
# NJ/PA/MD client's slug segment `princeton-nj` became the city "Princeton Nj".
# Same bug class, second site. These lock both directions: no client's home state
# is assumed, and Acme's behaviour is unchanged.

_NJ = {"states_served": ["NJ", "PA"]}
_NC = {"states_served": ["NC", "SC", "VA", "WV"]}


@pytest.mark.parametrize("slug,kind,config,expected", [
    # the regression: NJ suffix must be stripped for an NJ client
    ("north-nj/princeton-nj/demolition", "subservice", _NJ, "Princeton"),
    ("central-nj/trenton-nj",            "spoke",      _NJ, "Trenton"),
    # Acme unchanged
    ("charlotte-nc/matthews",            "spoke",      _NC, "Matthews"),
    ("charlotte-nc/mint-hill",           "spoke",      _NC, "Mint Hill"),
    # no config -> generic two-letter strip, never a hardcoded home state
    ("somewhere/gaithersburg-md",        "spoke",      None, "Gaithersburg"),
])
def test_city_from_slug_uses_client_states(slug, kind, config, expected):
    assert _city_from_slug(slug, kind, config) == expected


def test_city_from_slug_does_not_assume_nc_for_other_clients():
    """An NJ client must never inherit Acme's states via the default path."""
    assert _city_from_slug("north-nj/princeton-nj", "spoke", _NJ) == "Princeton"
    # and the old hardcoded set must not be silently applied to an NJ config
    assert not _city_from_slug("north-nj/princeton-nj", "spoke", _NJ).lower().endswith("nj")


# ── BUG-016, the class: ONE document, TWO clients ────────────────────────────
# Every emitter fixture used to be a Acme document distilled against the Acme
# repo, so client-specific hardcoding could not show itself: the "default" and the
# "correct" answer were the same string. Phase 2 wires this chain into CI for FIVE
# clients, so the guard that matters is CROSS-CLIENT, not cross-document — the same
# bytes distilled under two different client-configs must produce two different
# businesses, and neither may be a literal in the code.

def _state_neutral_page(tmp_path) -> str:
    """A minimal page whose slug carries NO state suffix, so `state` can only come
    from the client's own config. Deliberately generic copy: nothing in this
    document names a trade, a state, or a client."""
    d = docx.Document()
    d.add_paragraph('Page Title (56 chars): Service in Princeton')
    d.add_paragraph('Meta Description: Local crews handling residential and commercial work '
                    'with clear written estimates and code-compliant installation throughout.')
    d.add_paragraph('Canonical URL: https://example.com/north-region/princeton/site-work/')
    d.add_paragraph('HERO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('H1: Service in Princeton')
    d.add_paragraph('Hero Subheading:')
    d.add_paragraph('Our crews handle the whole job start to finish and leave the site clean '
                    'for every property owner we work with in the surrounding area.')
    d.add_paragraph('INTRO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Section Label: Local Crews')
    d.add_paragraph('H2: Work Built for Local Properties')
    d.add_paragraph('Paragraph 1: Every job starts with a walk of the property and a written '
                    'scope, so the number quoted is the number invoiced when the work ends.')
    return _save(d, tmp_path, 'state_neutral.docx')


_NJ_EXCAVATION = {"states_served": ["NJ", "PA"],
                  "trade": "demolition, excavation, dump truck rental"}
_NC_ROOFING = {"states_served": ["NC", "SC", "VA", "WV"],
               "trade": "roofing-contractor"}


def test_same_docx_two_clients_yields_two_businesses(tmp_path):
    """The same bytes, distilled under two client configs, must describe two
    different businesses in two different states — and neither answer may be
    reachable without a config."""
    docx_path = _state_neutral_page(tmp_path)

    nj = distill(docx_path, config=_NJ_EXCAVATION)[0]
    nc = distill(docx_path, config=_NC_ROOFING)[0]

    assert (nj.service, nj.state) == ('Demolition', 'NJ')
    assert (nc.service, nc.state) == ('Roofing Contractor', 'NC')
    # title = "<service> in <city>, <state>"; the city comes from the document
    # (identical for both runs), the service and the state come from the config.
    assert nj.title.startswith('Demolition in') and nj.title.endswith(', NJ')
    assert nc.title.startswith('Roofing Contractor in') and nc.title.endswith(', NC')
    # the two runs must genuinely differ — a hardcode makes them identical
    assert nj.service != nc.service and nj.state != nc.state


def test_missing_trade_or_states_is_a_blocking_finding_not_a_default(tmp_path):
    """No config = no invented business. BUG-016 was worse than a crash precisely
    because it produced confident, well-formed, wrong output; the contract now is
    fail-loud."""
    draft = distill(_state_neutral_page(tmp_path), config={})[0]
    codes = {f.code for f in draft.findings if f.blocking}
    assert 'no_service_label' in codes
    assert 'no_state' in codes
    assert draft.service == ''
    # and nothing anywhere silently claims the pilot's trade or geography
    assert 'Roofing' not in draft.title and 'NC' not in draft.title
