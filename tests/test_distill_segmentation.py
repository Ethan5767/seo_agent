"""Format-tolerant DOCX segmentation + the fail-loud guard.

Covers the two July-2026 defects fixed in pipeline.generate.distill:

  1. FORMAT-BRITTLE SEGMENTATION — segment_pages recognised only April's
     `Title`-style page boundaries, so July's 'Page Title:' / 'Canonical URL:' /
     'FULL PAGE CONTENT' handoff segmented to 0 pages.
  2. SILENT FALSE ALL-CLEAR — a non-empty DOCX that segments to 0 pages was
     treated as clean (exit 0). It must now FAIL LOUD with a distinct exit code.

Everything here is hermetic: DOCX files are synthesised under tmp_path with
python-docx; no real client repo or network is touched.
"""
from __future__ import annotations

import docx
import pytest

from pipeline.generate.distill import (
    UNSEGMENTABLE_EXIT,
    UnsegmentableDocxError,
    assert_segmentable,
    read_docx_blocks,
    segment_pages,
)


def _save(document: docx.document.Document, tmp_path, name: str) -> str:
    path = tmp_path / name
    document.save(str(path))
    return str(path)


def _july_doc(tmp_path) -> str:
    """A single July-format page: 'Page Title (NN chars):' boundary, a
    'Canonical URL:' identity line, and a 'FULL PAGE CONTENT' divider — none of
    the April `Title`-style markers."""
    d = docx.Document()
    d.add_paragraph('Page Title (56 chars): Soffit Installation in Matthews, NC | Since 2000')
    d.add_paragraph('Meta Description (155 chars): Professional soffit installation in '
                    'Matthews, NC. Vented, durable, code-compliant. Call today for a free '
                    'on-site estimate for your home project right now here today.')
    d.add_paragraph('Canonical URL: https://example.com/charlotte-nc/matthews/soffit-installation/')
    d.add_paragraph('FULL PAGE CONTENT', style=d.styles['Heading 3'])
    d.add_paragraph('HERO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('H1: Soffit Installation in Matthews, NC')
    d.add_paragraph('Hero Subheading: Soffit protects your roof from moisture and pests.')
    for i in range(1, 6):
        d.add_paragraph(f'Paragraph {i}: Body copy about soffit installation number {i}.')
    return _save(d, tmp_path, 'july.docx')


# ---------------------------------------------------------------------------
# July-format recognition
# ---------------------------------------------------------------------------

def test_july_page_title_convention_segments_to_one_page(tmp_path):
    blocks = read_docx_blocks(_july_doc(tmp_path))
    segments = segment_pages(blocks)
    assert len(segments) == 1
    # The boundary line is INSIDE the page so the meta/page title is reachable.
    page = segments[0]
    assert any(b.label.strip().lower() == 'page title' for b in page.blocks)
    # Canonical URL is normalised into the identity label set.
    assert any(b.label.strip().lower() == 'canonical url' for b in page.blocks)
    # A recognised page must NOT trip the fail-loud guard.
    assert_segmentable(blocks, segments, source='july.docx')


def test_full_page_content_backstop_when_no_page_title(tmp_path):
    """Convention 3: a 'FULL PAGE CONTENT' divider paired with an H1 + URL still
    segments even when the 'Page Title:' marker is absent."""
    d = docx.Document()
    d.add_paragraph('FULL PAGE CONTENT', style=d.styles['Heading 3'])
    d.add_paragraph('H1: Soffit Installation in Matthews, NC')
    d.add_paragraph('Canonical URL: https://example.com/charlotte-nc/matthews/')
    for i in range(1, 6):
        d.add_paragraph(f'Paragraph {i}: Body copy number {i}.')
    blocks = read_docx_blocks(_save(d, tmp_path, 'fpc.docx'))
    segments = segment_pages(blocks)
    assert len(segments) == 1
    assert segments[0].name == 'Soffit Installation in Matthews, NC'


# ---------------------------------------------------------------------------
# April regression — Title-style precedence (no double-split)
# ---------------------------------------------------------------------------

def test_title_style_convention_not_double_split_by_page_title(tmp_path):
    """April carries BOTH a `Title` block AND a 'Page Title:' line per page.
    Title-style must take precedence so each page stays whole (2 pages, not 4)."""
    d = docx.Document()
    for city in ('Charlotte', 'Matthews'):
        d.add_paragraph(city, style=d.styles['Title'])
        d.add_paragraph(f'Page Title: Roofing in {city}, NC | 5000+ Projects')
        d.add_paragraph(f'Canonical URL: https://example.com/{city.lower()}-nc/')
        d.add_paragraph(f'H1: Roofing in {city}')
    blocks = read_docx_blocks(_save(d, tmp_path, 'april.docx'))
    segments = segment_pages(blocks)
    assert [s.name for s in segments] == ['Charlotte', 'Matthews']


def test_week_dividers_still_dropped(tmp_path):
    d = docx.Document()
    d.add_paragraph('Week 1', style=d.styles['Title'])
    d.add_paragraph('Charlotte', style=d.styles['Title'])
    d.add_paragraph('Page Title: Roofing in Charlotte, NC')
    d.add_paragraph('H1: Roofing in Charlotte')
    blocks = read_docx_blocks(_save(d, tmp_path, 'week.docx'))
    segments = segment_pages(blocks)
    assert [s.name for s in segments] == ['Charlotte']


# ---------------------------------------------------------------------------
# Fail-loud guard
# ---------------------------------------------------------------------------

def test_unsegmentable_nonempty_docx_raises_with_exit_16(tmp_path):
    d = docx.Document()
    for i in range(25):
        d.add_paragraph(f'Ordinary prose line {i} with no page boundary marker anywhere.')
    blocks = read_docx_blocks(_save(d, tmp_path, 'blob.docx'))
    segments = segment_pages(blocks)
    assert segments == []
    with pytest.raises(UnsegmentableDocxError) as excinfo:
        assert_segmentable(blocks, segments, source='blob.docx')
    assert excinfo.value.exit_code == UNSEGMENTABLE_EXIT == 16
    msg = str(excinfo.value)
    assert '0 pages' in msg and 'unrecognized handoff format' in msg


def test_any_heading_alone_triggers_the_guard(tmp_path):
    """Below the paragraph threshold but carrying a Heading = still a parse
    failure, never a clean empty run."""
    d = docx.Document()
    d.add_paragraph('INTRO SECTION', style=d.styles['Heading 4'])
    d.add_paragraph('Some body copy with no recognizable page boundary.')
    blocks = read_docx_blocks(_save(d, tmp_path, 'heading.docx'))
    segments = segment_pages(blocks)
    assert segments == []
    with pytest.raises(UnsegmentableDocxError):
        assert_segmentable(blocks, segments, source='heading.docx')


def test_truly_empty_doc_does_not_raise(tmp_path):
    """A genuinely empty handoff (no meaningful content) returns quietly — there
    is nothing to parse, so the guard must not cry wolf."""
    d = docx.Document()
    d.add_paragraph('')
    d.add_paragraph('   ')
    blocks = read_docx_blocks(_save(d, tmp_path, 'empty.docx'))
    segments = segment_pages(blocks)
    assert segments == []
    assert_segmentable(blocks, segments, source='empty.docx')  # does not raise
